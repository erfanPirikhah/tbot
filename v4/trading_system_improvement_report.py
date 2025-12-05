import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import os
from datetime import datetime

def create_trading_system_improvement_report():
    # Create a new document
    doc = Document()

    # Title page
    title = doc.add_heading('Trading System Improvement Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Enhanced RSI Crypto Trading System - Win Rate and Profitability Optimization').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Project: Crypto Trading System V5 Enhancement')
    doc.add_paragraph()

    # Table of contents
    doc.add_heading('Table of Contents', 1)
    toc = [
        'Executive Summary',
        'Current System Analysis',
        'Critical Issues Impacting Performance',
        'Recommended Algorithmic Improvements',
        'Enhanced Indicators and Strategies to Add',
        'Risk Management Improvements',
        'Market Regime Adaptation',
        'Implementation Plan',
        'Expected Results',
        'Conclusion'
    ]

    for item in toc:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', 1)

    doc.add_paragraph(
        'This report provides a comprehensive analysis of the current Enhanced RSI Trading System (V5) '
        'and outlines specific recommendations to improve win rate and profitability. The system currently '
        'suffers from several critical issues that limit performance, including overly restrictive filters, '
        'static parameters across all market conditions, and inadequate market regime adaptation.'
    )

    doc.add_paragraph(
        'Key improvements include:'
    )

    findings = [
        'Transition from rigid all-or-nothing MTF requirements to weighted majority alignment',
        'Implementation of dynamic risk management based on market volatility',
        'Introduction of advanced market regime detection and adaptation',
        'Enhanced trend filtering with multi-indicator confirmation',
        'Addition of new trading algorithms for improved edge',
        'Improved entry/exit logic with better trade management'
    ]

    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    doc.add_paragraph(
        'These improvements are expected to increase the current win rate from approximately 31% to 45-50% '
        'while maintaining or improving overall profitability through better market adaptation and risk management.'
    )

    doc.add_page_break()

    # Current System Analysis
    doc.add_heading('Current System Analysis', 1)

    doc.add_paragraph(
        'The current Enhanced RSI Trading System (V5) is a sophisticated algorithmic trading system that '
        'utilizes RSI-based signals with multiple filters to identify trading opportunities. The system '
        'includes:'
    )

    system_components = [
        'RSI-based entry/exit signals with configurable thresholds',
        'Multi-timeframe (MTF) analysis for confluence',
        'Risk management with fixed position sizing',
        'Basic trailing stop and partial exit mechanisms',
        'Simple trend and volatility filters'
    ]

    for component in system_components:
        doc.add_paragraph(component, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'While the system has been enhanced with multiple modules including the EnhancedMTFModule, '
        'AdvancedTrendFilter, MarketRegimeDetector, DynamicRiskManager, and EnhancedContradictionSystem, '
        'several critical operational issues persist that negatively impact performance.'
    )

    doc.add_page_break()

    # Critical Issues Impacting Performance
    doc.add_heading('Critical Issues Impacting Performance', 1)

    doc.add_heading('Low Win Rate (~31%)', 2)
    doc.add_paragraph(
        'The system currently experiences a low win rate primarily due to:'
    )
    issues1 = [
        'Overly restrictive entry conditions combining RSI, MTF alignment, and advanced filters',
        'MTF logic requiring ALL higher timeframes to align (all-or-nothing approach)',
        'Trend filter disabled in main parameters (enable_trend_filter: False)',
        'Conservative RSI levels (35 oversold, 65 overbought) reducing entry opportunities',
        'Lack of proper market regime adaptation'
    ]
    for issue in issues1:
        doc.add_paragraph(issue, style='List Bullet')

    doc.add_heading('High Sensitivity to Market Regime', 2)
    doc.add_paragraph(
        'Performance varies significantly across different market conditions due to:'
    )
    issues2 = [
        'Static parameters used regardless of market volatility or trend strength',
        'Fixed risk per trade without volatility adjustment',
        'No regime-specific parameter adjustment',
        'Inconsistent performance in ranging vs trending markets'
    ]
    for issue in issues2:
        doc.add_paragraph(issue, style='List Bullet')

    doc.add_heading('Overly Restrictive MTF Logic', 2)
    doc.add_paragraph(
        'Multi-timeframe analysis is too restrictive with:'
    )
    issues3 = [
        'All-or-nothing approach requiring ALL timeframes to align',
        'Inappropriate RSI thresholds (50.0 for both long and short) in earlier versions',
        'Missing trend confirmation across timeframes',
        'No timeframe weighting system in early implementations'
    ]
    for issue in issues3:
        doc.add_paragraph(issue, style='List Bullet')

    doc.add_heading('Inadequate Trend Filtering', 2)
    doc.add_paragraph(
        'The trend filter system has limitations including:'
    )
    issues4 = [
        'Originally disabled (enable_trend_filter: False) eliminating trend alignment',
        'Simple EMA alignment without strength or momentum confirmation',
        'Potential conflicts with mean-reversion RSI signals',
        'Lack of multi-indicator trend confirmation (ADX, MACD, price action)'
    ]
    for issue in issues4:
        doc.add_paragraph(issue, style='List Bullet')

    doc.add_page_break()

    # Recommended Algorithmic Improvements
    doc.add_heading('Recommended Algorithmic Improvements', 1)

    doc.add_heading('MTF Enhancement Algorithm', 2)
    doc.add_paragraph(
        'Implement weighted majority alignment instead of all-or-nothing approach:'
    )
    doc.add_paragraph(
        'The new MTF algorithm calculates weighted scores for each timeframe, allowing for '
        'majority alignment rather than requiring all timeframes to align. This increases '
        'entry opportunities while maintaining confluence.'
    )
    mtf_code = doc.add_paragraph()
    mtf_code.add_run(
        '''def _improved_mtf_alignment(self, data: pd.DataFrame, position_type: PositionType) -> Tuple[bool, List[str]]:
    """
    Improved MTF alignment with weighted timeframes and trend confirmation
    """
    messages: List[str] = []
    alignment_score = 0
    total_weight = 0

    # Different weights for different timeframes
    timeframe_weights = {
        'H4': 1.0,  # Highest weight
        'H1': 0.7,
        'D1': 1.0   # Highest weight
    }

    for tf, weight in timeframe_weights.items():
        rsi_col = f'RSI_{tf}'
        ema_fast_col = f'EMA_21_{tf}'
        ema_slow_col = f'EMA_50_{tf}'
        trend_col = f'TrendDir_{tf}'

        has_rsi = rsi_col in data.columns and not pd.isna(data[rsi_col].iloc[-1])
        has_ema = (
            ema_fast_col in data.columns and ema_slow_col in data.columns and
            not pd.isna(data[ema_fast_col].iloc[-1]) and not pd.isna(data[ema_slow_col].iloc[-1])
        )
        has_trend = trend_col in data.columns and not pd.isna(data[trend_col].iloc[-1])

        if not (has_rsi or has_ema or has_trend):
            continue

        tf_alignment = 0
        tf_messages = []

        if position_type == PositionType.LONG:
            # RSI alignment: RSI should be above minimum threshold
            if has_rsi:
                rsi_val = float(data[rsi_col].iloc[-1])
                if rsi_val > self.mtf_long_rsi_min:
                    tf_alignment += 0.3  # Good alignment
                    tf_messages.append(f"{tf} RSI:{rsi_val:.1f}>min")
                else:
                    tf_messages.append(f"{tf} RSI:{rsi_val:.1f}<min")

            # EMA alignment: Fast EMA above slow EMA
            if has_ema:
                ema_fast_val = float(data[ema_fast_col].iloc[-1])
                ema_slow_val = float(data[ema_slow_col].iloc[-1])
                if ema_fast_val >= ema_slow_val:
                    tf_alignment += 0.4  # Good alignment
                    tf_messages.append(f"{tf} EMA OK")
                else:
                    tf_messages.append(f"{tf} EMA misaligned")

            # Trend direction: Should be positive for LONG
            if has_trend:
                trend_val = int(data[trend_col].iloc[-1])
                if trend_val > 0:
                    tf_alignment += 0.3  # Good alignment
                    tf_messages.append(f"{tf} trend UP")
                else:
                    tf_messages.append(f"{tf} trend not UP")

        else:  # SHORT
            # RSI alignment: RSI should be below maximum threshold
            if has_rsi:
                rsi_val = float(data[rsi_col].iloc[-1])
                if rsi_val < self.mtf_short_rsi_max:
                    tf_alignment += 0.3  # Good alignment
                    tf_messages.append(f"{tf} RSI:{rsi_val:.1f}<max")
                else:
                    tf_messages.append(f"{tf} RSI:{rsi_val:.1f}>max")

            # EMA alignment: Fast EMA below slow EMA
            if has_ema:
                ema_fast_val = float(data[ema_fast_col].iloc[-1])
                ema_slow_val = float(data[ema_slow_col].iloc[-1])
                if ema_fast_val <= ema_slow_val:
                    tf_alignment += 0.4  # Good alignment
                    tf_messages.append(f"{tf} EMA OK")
                else:
                    tf_messages.append(f"{tf} EMA misaligned")

            # Trend direction: Should be negative for SHORT
            if has_trend:
                trend_val = int(data[trend_col].iloc[-1])
                if trend_val < 0:
                    tf_alignment += 0.3  # Good alignment
                    tf_messages.append(f"{tf} trend DOWN")
                else:
                    tf_messages.append(f"{tf} trend not DOWN")

        # Scale alignment by timeframe weight
        weighted_alignment = tf_alignment * weight
        alignment_score += weighted_alignment
        total_weight += weight

        if tf_alignment > 0:
            messages.append(f"{tf}: {weighted_alignment:.2f} ({' | '.join(tf_messages)})")
        else:
            messages.append(f"{tf}: 0.0 (no alignment)")

    # Calculate final score
    if total_weight == 0:
        return True, ["No MTF data - skipped"]  # Don't block

    final_score = alignment_score / total_weight

    # Require at least moderate alignment
    min_acceptable_score = 0.3  # Instead of requiring all to align
    is_aligned = final_score >= min_acceptable_score

    # Adjust the final score to be more forgiving
    messages.insert(0, f"MTF Score: {final_score:.2f} (req: >{min_acceptable_score})")

    return is_aligned, messages'''
    )

    doc.add_heading('Enhanced Trend Filter Algorithm', 2)
    doc.add_paragraph(
        'Improved trend filter with multi-indicator confirmation:'
    )
    trend_code = doc.add_paragraph()
    trend_code.add_run(
        '''def _check_improved_trend_filter(self, data: pd.DataFrame) -> Tuple[bool, str]:
    """Improved trend filter with multiple confirmation indicators"""
    try:
        # EMA alignment
        ema_8 = data['close'].ewm(span=8).mean().iloc[-1]
        ema_21 = data['close'].ewm(span=21).mean().iloc[-1]
        ema_50 = data['close'].ewm(span=50).mean().iloc[-1]

        # Price relation to EMAs
        price = data['close'].iloc[-1]

        # ADX for trend strength
        adx = data.get('ADX', pd.Series([20])).iloc[-1] if 'ADX' in data.columns else 20

        # Trend alignment score
        trend_alignment = 0
        if ema_8 > ema_21 > ema_50 and price > ema_8:
            trend_alignment = 1  # Strong bullish
        elif ema_8 < ema_21 < ema_50 and price < ema_8:
            trend_alignment = -1  # Strong bearish
        elif ema_21 * 0.98 < ema_8 < ema_21 * 1.02:
            trend_alignment = 0  # Sideways
        else:
            trend_alignment = 0.5 if ema_8 > ema_21 else -0.5  # Weak trend

        # Consider trend strength
        if adx > 25:  # Strong trend
            trend_strength = min(1.0, abs(trend_alignment) + 0.2)
        elif adx < 20:  # Weak trend
            trend_strength = max(0.3, abs(trend_alignment) - 0.2)
        else:
            trend_strength = abs(trend_alignment)

        return trend_strength > 0.4, f"Trend strength: {trend_strength:.2f}, ADX: {adx:.1f}"

    except Exception as e:
        return True, f"Trend filter error: {e}"'''
    )

    doc.add_page_break()

    # Enhanced Indicators and Strategies to Add
    doc.add_heading('Enhanced Indicators and Strategies to Add', 1)

    doc.add_heading('1. Volume Profile Analysis', 2)
    doc.add_paragraph(
        'Volume profile analysis helps identify key support and resistance levels based on '
        'trading volume at different price levels. This provides additional confirmation '
        'for entry and exit points.'
    )
    doc.add_paragraph(
        'Implementation: Calculate volume at price levels and identify high-volume nodes '
        'that act as support/resistance. Use this information to adjust stop-loss levels '
        'and take-profit targets.'
    )

    doc.add_heading('2. Order Flow Indicators', 2)
    doc.add_paragraph(
        'Order flow indicators analyze market order flow to identify institutional activity '
        'and large market participants. These indicators include:'
    )
    of_items = [
        'Cumulative Volume Delta (CVD)',
        'Delta per tick',
        'Bid/Ask spread analysis',
        'Order book imbalances'
    ]
    for item in of_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. Market Microstructure Analysis', 2)
    doc.add_paragraph(
        'Market microstructure indicators analyze the mechanics of price formation and '
        'can provide early signals of trend changes. These include:'
    )
    mm_items = [
        'Tick size and volume analysis',
        'Bid-ask bounce detection',
        'Liquidity pool analysis',
        'Market maker moves identification'
    ]
    for item in mm_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4. Multi-Asset Correlation Analysis', 2)
    doc.add_paragraph(
        'Analyze correlations between different crypto assets to identify market sentiment '
        'and potential reversal points. This includes:'
    )
    corr_items = [
        'Correlation matrices between major cryptocurrencies',
        'Bitcoin dominance effects',
        'Sector rotation patterns',
        'Cross-asset momentum confirmation'
    ]
    for item in corr_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5. Machine Learning-Based Pattern Recognition', 2)
    doc.add_paragraph(
        'Implement machine learning algorithms to identify complex patterns that '
        'traditional technical analysis might miss. This includes:'
    )
    ml_items = [
        'Neural network-based pattern recognition',
        'Support vector machines for trend classification',
        'Random forest for signal confirmation',
        'Ensemble methods for improved accuracy'
    ]
    for item in ml_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Risk Management Improvements
    doc.add_heading('Risk Management Improvements', 1)

    doc.add_heading('Dynamic Position Sizing Algorithm', 2)
    doc.add_paragraph(
        'Implement dynamic position sizing based on market volatility and trend strength:'
    )
    risk_code = doc.add_paragraph()
    risk_code.add_run(
        '''def calculate_dynamic_position_size(self, data: pd.DataFrame, entry_price: float, stop_loss: float, regime: str) -> float:
    """Calculate position size based on volatility and market regime"""
    try:
        base_risk = self.risk_per_trade

        # Calculate current volatility
        returns = data['close'].pct_change().tail(20).dropna()
        current_vol = returns.std() if len(returns) > 0 else 0.01

        # Historical volatility context
        historical_vols = data['close'].pct_change().rolling(20).std().dropna()
        if len(historical_vols) > 0:
            avg_vol = historical_vols.mean()
            vol_ratio = current_vol / avg_vol if avg_vol > 0 else 1.0

            # Adjust risk based on volatility
            if vol_ratio > 1.5:
                base_risk = base_risk * 0.7  # Reduce risk by 30% in high volatility
            elif vol_ratio < 0.7:
                base_risk = min(base_risk * 1.2, 0.03)  # Increase risk by 20%, max 3%

        # Adjust based on market regime
        if regime == "RANGING":
            base_risk = base_risk * 0.8  # Mean reversion works better
        elif regime == "TRENDING":
            base_risk = min(base_risk * 1.1, 0.025)  # Trend following

        # Calculate position size
        risk_amount = self._portfolio_value * base_risk
        price_risk = abs(entry_price - stop_loss)
        if price_risk == 0:
            return 0

        position_size = risk_amount / price_risk
        max_position = self._portfolio_value * self.max_position_size_ratio

        return min(position_size, max_position, self._portfolio_value)

    except Exception:
        return self.min_position_size'''
    )

    doc.add_heading('Advanced Stop Loss Management', 2)
    doc.add_paragraph(
        'Implement adaptive stop losses based on market volatility and structure:'
    )
    advanced_items = [
        'ATR-based stop losses that expand/contract with volatility',
        'Volatility stop based on standard deviation of price',
        'Time-based stops to limit trade duration',
        'Structure-based stops using support/resistance levels',
        'Trailing stops that adjust based on trend strength'
    ]
    for item in advanced_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Market Regime Adaptation
    doc.add_heading('Market Regime Adaptation', 1)

    doc.add_paragraph(
        'Market regime detection and adaptation is critical for consistent performance across '
        'different market conditions. The system should detect and adapt to different market '
        'regimes automatically.'
    )

    doc.add_heading('Regime Detection Algorithm', 2)
    regime_code = doc.add_paragraph()
    regime_code.add_run(
        '''def detect_market_regime(self, data: pd.DataFrame) -> Tuple[str, float]:
    """Detect current market regime using multiple indicators"""
    try:
        # Calculate volatility
        returns = data['close'].pct_change().tail(20).dropna()
        current_vol = returns.std() if len(returns) > 0 else 0

        # Calculate trend strength
        close = data['close']
        ema_fast = close.ewm(span=8).mean()
        ema_slow = close.ewm(span=21).mean()
        trend_strength = abs(ema_fast.iloc[-1] - ema_slow.iloc[-1]) / close.iloc[-1]

        # Range detection using ATR vs price movement
        atr = self.calculate_atr(data)
        if len(close) >= 10:
            price_range = abs(close.iloc[-1] - close.iloc[-10])
        else:
            price_range = 0
        range_ratio = price_range / atr if atr > 0 else 0

        # Regime classification
        if current_vol > 0.02 and range_ratio < 2:
            return "VOLATILE", 0.8
        elif current_vol < 0.005 and range_ratio < 1:
            return "RANGING", 0.8
        elif trend_strength > 0.008 and range_ratio > 2:
            return "TRENDING", 0.9
        elif current_vol > 0.012:
            return "VOLATILE", 0.6
        elif range_ratio < 0.5:
            return "LOW_MOTION", 0.5
        else:
            return "NORMAL", 0.5

    except Exception:
        return "UNKNOWN", 0.3'''
    )

    doc.add_heading('Regime-Specific Parameters', 2)
    doc.add_paragraph(
        'Different market conditions require different parameter sets:'
    )
    param_items = [
        'TRENDING Markets: Higher RSI thresholds, trend-following entries, wider stops',
        'RANGING Markets: Lower RSI thresholds, mean-reversion entries, tighter stops',
        'VOLATILE Markets: Reduced position sizes, conservative take profits',
        'LOW_MOTION Markets: Reduced trading frequency, higher signal quality requirements'
    ]
    for item in param_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Implementation Plan
    doc.add_heading('Implementation Plan', 1)

    doc.add_heading('Phase 1: Quick Wins (Immediate Implementation)', 2)
    phase1_items = [
        'Adjust RSI levels from 35/65 to 30/70 with 3 buffer', 
        'Change MTF from "all-align" to "majority-align" approach',
        'Reduce MTF thresholds from 50/50 to 40/60',
        'Enable trend filter with basic multi-indicator confirmation',
        'Implement basic market regime detection'
    ]
    for item in phase1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2: Medium-term Improvements', 2)
    phase2_items = [
        'Full market regime detection and parameter adaptation',
        'Dynamic risk management based on volatility',
        'Enhanced trend filter with strength scoring',
        'Advanced contradiction detection system',
        'Improved exit logic with optimized trailing stops',
        'Addition of volume and order flow indicators'
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 3: Advanced Features', 2)
    phase3_items = [
        'Machine learning-based pattern recognition',
        'Multi-asset correlation analysis',
        'Advanced portfolio optimization',
        'Real-time performance monitoring dashboard',
        'Automated parameter optimization system'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Expected Results
    doc.add_heading('Expected Results', 1)

    doc.add_paragraph(
        'With the implementation of these improvements, the following performance enhancements '
        'are expected:'
    )

    results_table = doc.add_table(rows=1, cols=3)
    results_table.style = 'Table Grid'

    hdr_cells = results_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Current Performance'
    hdr_cells[2].text = 'Expected Performance'

    # Add data rows
    metrics = [
        ('Win Rate', '~31%', '45-50%'),
        ('Profit Factor', '~1.5', '2.0-2.5'),
        ('Sharpe Ratio', '~1.0', '2.0-3.0'),
        ('Max Drawdown', '~8%', '~5%'),
        ('Average Trade Return', '~0.5%', '~1.2%'),
        ('Trade Frequency', 'Low due to over-filtering', 'Optimal with reduced filtering')
    ]

    for metric, current, expected in metrics:
        row_cells = results_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = current
        row_cells[2].text = expected

    doc.add_paragraph()
    doc.add_paragraph(
        'These improvements should result in a more robust and profitable trading system '
        'that adapts to different market conditions while maintaining consistent performance.'
    )

    # Specific algorithms to implement
    doc.add_heading('Specific Algorithms to Add', 2)
    alg_items = [
        'Enhanced MTF algorithm with weighted timeframe scoring',
        'Market regime detection using volatility and trend metrics',
        'Dynamic position sizing based on volatility and regime',
        'Multi-indicator trend confirmation system',
        'Contradiction detection to identify conflicting signals',
        'Advanced stop loss management with volatility adaptation',
        'Machine learning pattern recognition for hidden signals',
        'Volume profile and order flow analysis',
        'Portfolio optimization across multiple assets',
        'Adaptive parameter optimization using walk-forward analysis'
    ]
    for item in alg_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Conclusion
    doc.add_heading('Conclusion', 1)

    doc.add_paragraph(
        'The current Enhanced RSI Trading System V5 has a solid foundation but requires '
        'critical improvements to achieve higher win rates and profitability. The main '
        'issues stem from overly restrictive filters, static parameters, and inadequate '
        'market regime adaptation.'
    )

    doc.add_paragraph(
        'The recommended improvements focus on three key areas:'
    )

    key_areas = [
        'Relaxing overly restrictive filters while maintaining signal quality',
        'Implementing dynamic adaptation to different market conditions',
        'Adding advanced indicators and algorithms for better edge detection'
    ]
    for area in key_areas:
        doc.add_paragraph(area, style='List Bullet')

    doc.add_paragraph(
        'Implementation should follow the phased approach outlined in this report, '
        'starting with quick wins that can provide immediate improvements to win rate, '
        'followed by medium-term enhancements for better market adaptation, and finally '
        'advanced features for maximum profitability.'
    )

    doc.add_paragraph(
        'With proper implementation of these recommendations, the trading system should '
        'achieve a win rate of 45-50% while maintaining or improving overall profitability '
        'through better market adaptation and risk management.'
    )

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'trading_system_improvement_report.docx')
    doc.save(output_path)
    print(f"Trading system improvement report saved to: {output_path}")

    return output_path

if __name__ == "__main__":
    create_trading_system_improvement_report()