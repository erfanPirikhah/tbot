import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_document_style(doc: Document, font_name: str = "Tahoma", font_size: int = 11) -> None:
    """
    Configure default font for Persian text rendering.
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    # Ensure complex script font is set for Persian
    rpr = style.element.rPr
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_content(doc: Document) -> None:
    # Title
    add_title(doc, "مقایسه استراتژی‌های v4 با چارچوب Freqtrade")

    # Executive Summary
    add_heading(doc, "خلاصه اجرایی", level=1)
    add_paragraph(
        doc,
        "چارچوب Freqtrade یک اکوسیستم حرفه‌ای برای توسعه، بک‌تست، هایپراپت، استقرار و مانیتورینگ استراتژی‌های معاملاتی است. "
        "استراتژی‌های v4 شما (EnhancedRsiStrategyV4 و EnsembleRsiStrategyV4) از نظر منطق سیگنال‌دهی، مدیریت ریسک بر پایه ATR، خروج‌های مرحله‌ای "
        "و کنترل‌های عملیاتی (فاصله بین معاملات، توقف پس از ضرر) بسیار کامل و مناسب برای بازارهای اینترادی هستند.",
    )
    add_bullets(
        doc,
        [
            "Freqtrade برای عملیات حرفه‌ای (Backtesting, Hyperopt, REST/Web UI, Docker) برتر است.",
            "استراتژی‌های v4 برای منطق معاملاتی سفارشی و انعطاف بالا در تایم‌فریم‌های M5/M15/H1 بسیار مؤثرند.",
            "جمع‌بندی: برای «زیرساخت حرفه‌ای»، Freqtrade بهتر است؛ برای «منطق سیگنال‌گیری سفارشی»، v4 شما بالغ و حرفه‌ای است.",
        ],
    )

    # Criteria
    add_heading(doc, "معیارهای مقایسه", level=1)
    add_bullets(
        doc,
        [
            "معماری و طراحی استراتژی",
            "ویژگی‌های سیگنال و فیلترها",
            "مدیریت ریسک و خروج‌ها",
            "بهینه‌سازی / Hyperopt / ML",
            "قابلیت توسعه و نگهداشت",
            "ابزارهای عملیاتی (UI, API, گزارش‌گیری)",
        ],
    )

    # v4 strategies intro
    add_heading(doc, "معرفی استراتژی‌های v4", level=1)

    add_heading(doc, "EnhancedRsiStrategyV4 (نسخه بهینه‌شده RSI)", level=2)
    add_bullets(
        doc,
        [
            "درگاه RSI با بافر ورود (Oversold/Overbought + Entry Buffer).",
            "فیلتر روند اختیاری (همسویی EMA21/EMA50 و آستانه ADX).",
            "تأیید شیب RSI و تأیید قیمتی اختیاری (EMA9 و جهت کندل).",
            "مدیریت ریسک مبتنی بر ATR: محاسبه Stop Loss/Take Profit پویا.",
            "حداکثر معاملات در 100 کندل، حداقل فاصله بین معاملات، توقف پس از ضررهای متوالی.",
            "Trailing Stop پویا با فعال‌سازی درصد سود و ضریب ATR.",
            "خروج مرحله‌ای (Partial Exit) برای قفل کردن سود.",
            "MTF اختیاری با بررسی هم‌راستایی تایم‌فریم‌های بالاتر.",
        ],
    )

    add_heading(doc, "EnsembleRsiStrategyV4 (ترکیبی برای اسکالپینگ)", level=2)
    add_bullets(
        doc,
        [
            "ترکیب چند میکرو‌سیگنال: Mean Reversion، Trend Pullback، Breakout، MACD+RSI.",
            "امتیازدهی وزنی و آستانه ورود با نسبت تسلط جهت غالب.",
            "فیلترهای جلسه معاملاتی و پهنای باند بولینگر برای اجتناب از بازارهای مرده یا اسپایک.",
            "مدیریت ریسک و خروج‌های ATR، خروج مرحله‌ای، Trailing Stop.",
            "کنترل‌های عملیاتی اسکالپینگ: فاصله کندلی کم، سقف معاملات در پنجره 100 کندل.",
        ],
    )

    # Freqtrade intro
    add_heading(doc, "چارچوب Freqtrade و نمونه استراتژی", level=1)
    add_bullets(
        doc,
        [
            "رابط IStrategy با چرخه کامل: populate_indicators، populate_entry_trend، populate_exit_trend، callbacks سفارشی.",
            "پشتیبانی کامل از هایپراپت (Hyperopt) با فضای پارامترهای قابل‌تعریف و توزیع‌ها (Int/Real/Decimal/Categorical).",
            "مدیریت ROI، Stoploss، Trailing، و منطق خروج سفارشی (custom_exit/custom_stoploss).",
            "ابزارهای داده، MTF و informative decorators، ادغام معاملات عمومی، کاهش footprint دیتافریم.",
            "زیرساخت‌های عملیاتی: Backtesting/Hyperopt، REST API، Web UI، Docker Compose، پشتیبانی چند صرافی.",
        ],
    )

    # Detailed comparison
    add_heading(doc, "مقایسه جزئی", level=1)

    add_heading(doc, "۱) معماری و طراحی", level=2)
    add_bullets(
        doc,
        [
            "v4: کلاس‌های سفارشی با حالت داخلی (Stateful) و API ساده generate_signal؛ مناسب برای سناریوهای اختصاصی.",
            "Freqtrade: IStrategy با قراردادهای روشن، اکوسیستم ابزار کامل و جداسازی concerns؛ مناسب برای تولید و استقرار.",
        ],
    )

    add_heading(doc, "۲) سیگنال‌ها و فیلترها", level=2)
    add_bullets(
        doc,
        [
            "v4: درگاه RSI، فیلتر روند EMA+ADX، تأیید قیمت/شیب، MTF اختیاری؛ Ensemble با امتیازدهی چند مؤلفه‌ای.",
            "Freqtrade: تعریف اندیکاتورها و سیگنال‌ها در DataFrame؛ کراس‌ها و گاردها (نمونه RSI+TEMA+BB).",
        ],
    )

    add_heading(doc, "۳) مدیریت ریسک و خروج‌ها", level=2)
    add_bullets(
        doc,
        [
            "v4: SL/TP بر اساس ATR، Trailing Stop با فعال‌سازی درصدی، خروج مرحله‌ای، محدودیت اندازه پوزیشن، توقف پس از ضرر.",
            "Freqtrade: ROI پله‌ای، Stoploss ثابت/سفارشی، Trailing مثبت با offset، custom_exit برای منطق پیشرفته.",
        ],
    )

    add_heading(doc, "۴) بهینه‌سازی و ML", level=2)
    add_bullets(
        doc,
        [
            "v4: بهینه‌سازی شبکه‌ای کوچک داخلی و گزارش‌گیری؛ ساده و سریع برای تیونینگ اولیه.",
            "Freqtrade: هایپراپت کامل با فضای پارامترها، معیارهای متعدد زیان/کارایی، FreqAI برای ویژگی‌سازی و مدل‌سازی ML.",
        ],
    )

    add_heading(doc, "۵) توسعه‌پذیری و نگهداشت", level=2)
    add_bullets(
        doc,
        [
            "v4: کد فشرده و خوانا، مناسب برای توسعه سریع و سفارشی‌سازی دقیق در بازارهای خاص.",
            "Freqtrade: استانداردسازی روش‌ها، تحلیل و اعتبارسنجی نتایج، ادغام آسان با UI/API و زیرساخت‌های مدرن.",
        ],
    )

    add_heading(doc, "۶) ابزارهای عملیاتی", level=2)
    add_bullets(
        doc,
        [
            "v4: گزارش JSON/PNG سفارشی، لاگ‌های دقیق؛ بدون UI وب داخلی.",
            "Freqtrade: Web UI، REST API، RPC (تلگرام/دیسکورد/وب‌هوک)، کانفیگ‌پذیری بالا، Docker و مقیاس‌پذیری.",
        ],
    )

    # Conclusion
    add_heading(doc, "نتیجه‌گیری", level=1)
    add_paragraph(
        doc,
        "اگر هدف شما زیرساخت حرفه‌ای با قابلیت مانیتورینگ، هایپراپت، استقرار و چند‌صرافی است، Freqtrade انتخاب بهتر و حرفه‌ای‌تر است. "
        "اگر هدف شما منطق سیگنال‌گیری سفارشی برای تایم‌فریم‌های اینترادی (مانند M5/M15/H1) با کنترل‌های دقیق ریسک و رفتار معاملاتی است، "
        "استراتژی‌های v4 شما بسیار بالغ و کارآمد هستند.",
        bold=True,
    )

    # Recommendations
    add_heading(doc, "توصیه‌ها", level=1)
    add_bullets(
        doc,
        [
            "پورت کردن EnhancedRsiStrategyV4 به Freqtrade (IStrategy) برای بهره‌مندی از هایپراپت و UI.",
            "استفاده از informative decorators برای پیاده‌سازی MTF در سبک Freqtrade.",
            "تعریف پارامترهای هایپراپت (Int/Real/Decimal/Boolean) برای بافر RSI، ضرایب ATR، آستانه‌های ADX و فاصله کندلی.",
            "ادغام خروج مرحله‌ای در custom_exit یا با مدیریت اندازه پوزیشن در Freqtrade.",
        ],
    )

    # References section (textual, not live links inside docx)
    add_heading(doc, "ارجاعات کد (برای توسعه‌دهندگان)", level=1)
    add_bullets(
        doc,
        [
            "EnhancedRsiStrategyV4: check_entry_conditions، مدیریت ATR SL/TP، Trailing، Partial Exit.",
            "EnsembleRsiStrategyV4: امتیازدهی Mean Reversion / Trend Pullback / Breakout / MACD+RSI.",
            "Freqtrade IStrategy: populate_indicators، populate_entry_trend، populate_exit_trend، custom_* callbacks.",
            "Freqtrade SampleStrategy: نمونه RSI/TEMA/BB با پارامترهای هایپراپت.",
        ],
    )


def build_docx(output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    set_document_style(doc)
    build_content(doc)
    doc.save(output_path)
    return output_path


def main():
    out_path = os.path.join("reports", "strategy_comparison_v4_vs_freqtrade.docx")
    final_path = build_docx(out_path)
    print(f"Saved report to: {final_path}")


if __name__ == "__main__":
    main()